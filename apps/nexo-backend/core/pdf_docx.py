# core/pdf_docx.py
# ─── Generación de PDF y DOCX para escritos legales ───────────────────────────
import io
import base64
from loguru import logger

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_RIGHT, TA_CENTER, TA_JUSTIFY
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False

try:
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False


def generar_pdf_basico(texto_escrito: str, nombre_usuario: str, rut_usuario: str) -> str:
    """Genera el escrito legal como PDF base64 con márgenes judiciales chilenos."""
    if not REPORTLAB_OK:
        return ""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        rightMargin=2.0 * cm, leftMargin=3.0 * cm,
        topMargin=2.5 * cm, bottomMargin=2.5 * cm,
    )
    styles = getSampleStyleSheet()
    style_normal = ParagraphStyle(
        'LegalNormal', parent=styles['Normal'],
        fontName='Helvetica', fontSize=12, leading=18, alignment=TA_JUSTIFY,
    )
    style_header = ParagraphStyle(
        'LegalHeader', parent=styles['Normal'],
        fontName='Helvetica-Bold', fontSize=12, leading=15, alignment=TA_RIGHT,
    )
    style_center = ParagraphStyle(name='Center', fontName='Helvetica', fontSize=12, alignment=TA_CENTER)
    style_small  = ParagraphStyle(name='Small',  fontName='Helvetica', fontSize=8,  alignment=TA_CENTER)

    story = []
    body_started = False

    for line in texto_escrito.split('\n'):
        text = line.strip()
        if not text:
            story.append(Spacer(1, 12))
            continue

        text_no_sp = text.upper().replace(" ", "")

        if not body_started and any(x in text_no_sp for x in ["SUMA:", "RIT:", "MATERIA:", "ENLOPRINCIPAL", "SOLICITA:"]):
            if "S.J.L." in text_no_sp or "S.S." in text_no_sp:
                body_started = True
                story.append(Spacer(1, 12))
                story.append(Paragraph(text, style_normal))
            else:
                story.append(Paragraph(text, style_header))
            continue

        if "S.J.L." in text.upper() or "S.S." in text.upper() or "S.S.ª" in text.upper():
            body_started = True
            story.append(Spacer(1, 12))

        if text.upper().startswith("OTROSÍ"):
            story.append(Paragraph(f"<b>{text}</b>", style_normal))
        else:
            story.append(Paragraph(text, style_normal))
        story.append(Spacer(1, 6))

    # Bloque firma
    story.append(Spacer(1, 3 * cm))
    story.append(Paragraph("___________________________", style_center))
    story.append(Paragraph(nombre_usuario.upper(), style_center))
    story.append(Paragraph(f"RUT: {rut_usuario}", style_center))
    story.append(Spacer(1, 1 * cm))
    story.append(Paragraph(
        "Documento generado para ser firmado mediante Clave Única en la Oficina Judicial Virtual.",
        style_small,
    ))

    try:
        doc.build(story)
        return base64.b64encode(buffer.getvalue()).decode('utf-8')
    except Exception as e:
        logger.error(f"[ReportLab] Error: {e}")
        return ""


def generar_docx_basico(texto_escrito: str, nombre_usuario: str, rut_usuario: str) -> str:
    """Genera el escrito legal como DOCX base64 con márgenes judiciales chilenos."""
    if not DOCX_OK:
        return ""
    try:
        doc = Document()
        for section in doc.sections:
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin   = Cm(3.0)
            section.right_margin  = Cm(2.0)

        body_started = False

        for line in texto_escrito.split('\n'):
            text = line.strip()
            if not text:
                continue

            text_no_sp = text.upper().replace(" ", "")
            p   = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(12)

            if not body_started and any(x in text_no_sp for x in ["SUMA:", "RIT:", "MATERIA:", "ENLOPRINCIPAL", "SOLICITA:"]):
                if "S.J.L." in text_no_sp or "S.S." in text_no_sp:
                    body_started = True
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    run.bold = True
                continue

            if "S.J.L." in text.upper() or "S.S." in text.upper() or "S.S.ª" in text.upper():
                body_started = True

            if text.upper().startswith("OTROSÍ"):
                run.bold    = True
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Firma
        doc.add_paragraph()
        p_firma = doc.add_paragraph()
        p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_f = p_firma.add_run(f"___________________________\n{nombre_usuario.upper()}\n[RUT: {rut_usuario}]")
        run_f.font.name = 'Arial'
        run_f.font.size = Pt(12)

        p_nota = doc.add_paragraph()
        p_nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_nota = p_nota.add_run("\nDocumento generado para ser firmado mediante Clave Única en la Oficina Judicial Virtual.")
        run_nota.font.name = 'Arial'
        run_nota.font.size = Pt(8)

        buffer = io.BytesIO()
        doc.save(buffer)
        return base64.b64encode(buffer.getvalue()).decode('utf-8')
    except Exception as e:
        logger.error(f"[python-docx] Error: {e}")
        return ""
