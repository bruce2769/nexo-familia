from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Depends, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel
import fitz  # PyMuPDF
import re
import sqlite3
import os
import json
import hashlib
import openai
import base64
import time
from collections import defaultdict
from datetime import date, datetime, timezone
import io

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import cm
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.enums import TA_RIGHT, TA_CENTER, TA_JUSTIFY
    REPORTLAB_OK = True
except ImportError:
    REPORTLAB_OK = False

try:
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

from loguru import logger
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded

# ─── Firebase Admin SDK ────────────────────────────────────────────────────────
try:
    import firebase_admin
    from firebase_admin import credentials, auth as firebase_auth, firestore

    _sa_json = os.environ.get("FIREBASE_SERVICE_ACCOUNT_JSON", "")
    if _sa_json and not firebase_admin._apps:
        try:
            # ── Parser robusto: Railway a veces almacena \n como \\n en la private_key ──
            # Intento 1: parse directo
            cred_dict = None
            for attempt, json_str in enumerate([
                _sa_json,                         # 1. Tal como está
                _sa_json.replace('\\n', '\n'),     # 2. Convertir \\n → \n (fix Railway)
                _sa_json.replace('\\n', '\n').replace("\\'", "'"),  # 3. Fix extra escapes
            ]):
                try:
                    cred_dict = json.loads(json_str, strict=False)
                    if attempt > 0:
                        logger.info(f"✅ Firebase JSON parseado en intento {attempt + 1} (fix \\n aplicado).")
                    break
                except json.JSONDecodeError:
                    continue

            if cred_dict is None:
                raise ValueError("No se pudo parsear FIREBASE_SERVICE_ACCOUNT_JSON después de 3 intentos.")

            cred = credentials.Certificate(cred_dict)
            firebase_admin.initialize_app(cred)
            db = firestore.client()
            FIREBASE_ADMIN_OK = True
            logger.info("✅ Firebase Admin SDK inicializado correctamente.")
        except Exception as e:
            logger.warning(f"⚠️ Error inicializando Firebase: {e}")
            FIREBASE_ADMIN_OK = False
            db = None
    else:
        if firebase_admin._apps:
            db = firestore.client()
            FIREBASE_ADMIN_OK = True
        else:
            FIREBASE_ADMIN_OK = False
            db = None
except ImportError:
    FIREBASE_ADMIN_OK = False
    db = None


# ─── Validación de variables de entorno (warning, no crash) ─────────────────────
REQUIRED_ENV = ["ALLOWED_ORIGINS", "OPENAI_API_KEY"]
missing_env = [k for k in REQUIRED_ENV if not os.environ.get(k)]
if missing_env:
    logger.warning(f"⚠️ Variables de entorno faltantes (modo degradado): {', '.join(missing_env)}")

_bearer_scheme = HTTPBearer(auto_error=False)

async def verify_firebase_token(
    request: Request,
    credentials: HTTPAuthorizationCredentials = Depends(_bearer_scheme)
):
    if not FIREBASE_ADMIN_OK:
        return {"uid": "test-user"}
    if credentials is None:
        logger.warning("Modo testing sin auth: Token ausente, asignando test-user.")
        return {"uid": "test-user"}
    try:
        decoded = firebase_auth.verify_id_token(credentials.credentials)
        return decoded
    except Exception as e:
        logger.warning(f"Modo testing sin auth: Token inválido '{str(e)[:80]}', asignando test-user.")
        return {"uid": "test-user"}

# ─── Configuración OpenAI ──────────────────────────────────────────────────────
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY")
OPENAI_MODEL   = os.environ.get("OPENAI_MODEL", "gpt-4o-mini")

# ─── Límite diario de uso de IA por usuario ────────────────────────────────────
MAX_IA_REQUESTS_PER_DAY = int(os.environ.get("MAX_IA_REQUESTS_PER_DAY", "20"))

# ─── Rate Limiter simple ───────────────────────────────────────────────────────
RATE_LIMITS = defaultdict(list)
MAX_REQUESTS_PER_MINUTE = 20

def check_rate_limit(client_ip: str):
    now = time.time()
    RATE_LIMITS[client_ip] = [t for t in RATE_LIMITS[client_ip] if now - t < 60]
    if len(RATE_LIMITS[client_ip]) >= MAX_REQUESTS_PER_MINUTE:
        raise HTTPException(status_code=429, detail="Too many requests. Please wait a minute.")
    RATE_LIMITS[client_ip].append(now)

# ─── FastAPI App ───────────────────────────────────────────────────────────────
app = FastAPI(
    title="Nexo Familia API",
    description="Motor backend para Legal Analytics — powered by OpenAI gpt-4o-mini.",
    version="2.0.0"
)

limiter = Limiter(key_func=get_remote_address, default_limits=["100/minute"])
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# ─── CORS global (permite cualquier origen: Vercel, localhost, etc.) ────────────
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ─── Middleware de logging de requests ─────────────────────────────────────────
@app.middleware("http")
async def log_requests(request: Request, call_next):
    logger.info(f"→ {request.method} {request.url.path} | client={request.client.host if request.client else 'unknown'}")
    response = await call_next(request)
    logger.info(f"← {request.method} {request.url.path} | status={response.status_code}")
    return response

# ─── Health check ──────────────────────────────────────────────────────────────
@app.get("/health", tags=["Health"])
async def health_check():
    return {"ok": True, "status": "ok", "version": "2.0.0", "ia": "openai", "model": OPENAI_MODEL}

@app.get("/api/health", tags=["Health"])
async def api_health_check():
    """Alias de /health para compatibilidad con frontends."""
    return {"ok": True, "status": "ok", "version": "2.0.0", "ia": "openai", "model": OPENAI_MODEL}

# ─── Helpers de base de datos SQLite ──────────────────────────────────────────
def get_db_connection():
    db_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'causas_judiciales.db')
    conn = sqlite3.connect(db_path)
    conn.row_factory = sqlite3.Row
    return conn

# ─── Utilidades: Privacidad y Caché MD5 ───────────────────────────────────────
def pseudo_anonimizar(texto: str) -> str:
    patron_rut = r'\b\d{1,2}\.?\d{3}\.?\d{3}[-][0-9kK]\b'
    return re.sub(patron_rut, "[RUT_OCULTO]", texto)

def validar_rut_chileno(rut: str) -> bool:
    if not rut: return False
    rut_limpio = rut.upper().replace("-", "").replace(".", "").strip()
    if len(rut_limpio) < 2: return False
    cuerpo, dv = rut_limpio[:-1], rut_limpio[-1]
    if not cuerpo.isdigit(): return False
    suma = 0
    multiplo = 2
    for c in reversed(cuerpo):
        suma += int(c) * multiplo
        multiplo = multiplo + 1 if multiplo < 7 else 2
    esperado = 11 - (suma % 11)
    if esperado == 11: dv_esp = '0'
    elif esperado == 10: dv_esp = 'K'
    else: dv_esp = str(esperado)
    return dv == dv_esp

def _generar_hash_md5(texto: str) -> str:
    texto_normalizado = re.sub(r'\s+', ' ', texto.strip().lower())
    return hashlib.md5(texto_normalizado.encode('utf-8')).hexdigest()

def _buscar_en_cache(hash_key: str):
    if not FIREBASE_ADMIN_OK or db is None:
        return None
    try:
        doc = db.collection("cache_ia").document(hash_key).get()
        if doc.exists:
            logger.info(f"[Cache] ✅ HIT {hash_key[:8]}... → $0")
            return doc.to_dict().get("respuesta")
    except Exception as e:
        logger.warning(f"[Cache] ⚠️ Error: {e}")
    return None

def _guardar_en_cache(hash_key: str, texto_preview: str, respuesta: dict):
    if not FIREBASE_ADMIN_OK or db is None:
        return
    try:
        db.collection("cache_ia").document(hash_key).set({
            "hash": hash_key,
            "texto_preview": texto_preview[:100],
            "respuesta": respuesta,
            "timestamp": firestore.SERVER_TIMESTAMP,
        })
    except Exception as e:
        logger.warning(f"[Cache] ⚠️ Error guardando: {e}")

def _verificar_limite_diario(uid: str) -> bool:
    if not FIREBASE_ADMIN_OK or db is None or not uid:
        return True
    try:
        hoy = date.today().isoformat()
        doc_ref = db.collection("usuarios").document(uid).collection("ia_usage").document(hoy)
        doc = doc_ref.get()
        if doc.exists:
            conteo = doc.to_dict().get("requests", 0)
            if conteo >= MAX_IA_REQUESTS_PER_DAY:
                return False
            doc_ref.update({"requests": conteo + 1})
        else:
            doc_ref.set({"requests": 1, "fecha": hoy})
    except Exception as e:
        logger.warning(f"[RateLimit] Error: {e}")
    return True

# ─── Motor IA: OpenAI gpt-4o-mini (ÚNICO motor) ───────────────────────────────
def llamar_openai(prompt: str, system: str = "", temperatura: float = 0.2, max_tokens: int = 200) -> str:
    """
    Llama a OpenAI gpt-4o-mini con parámetros optimizados.
    """
    try:
        client = openai.OpenAI(api_key=OPENAI_API_KEY)
        messages = []
        if system:
            messages.append({"role": "system", "content": system})
        messages.append({"role": "user", "content": prompt})

        response = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=messages,
            temperature=temperatura,
            max_tokens=max_tokens,
        )
        logger.info(f"[OpenAI] ✅ {response.usage.total_tokens} tokens usados.")
        return response.choices[0].message.content or ""
    except Exception as e:
        logger.error(f"[OpenAI] Error: {e}")
        return "{}"

def llamar_openai_vision(base64_image: str) -> str:
    """Llama a la visión de gpt-4o-mini para extraer transcripción perfecta."""
    try:
        client = openai.OpenAI(api_key=OPENAI_API_KEY)
        response = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Transcribe exactamente todo el texto de la imagen proporcionada de un documento legal. No añadas encabezados ni comentarios extra, solo el texto extraído tal cual."},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
                    ]
                }
            ],
            max_tokens=2000,
            temperature=0.1
        )
        logger.info(f"[OpenAI Vision OCR] ✅ {response.usage.total_tokens} tokens usados.")
        return response.choices[0].message.content or ""
    except Exception as e:
        logger.error(f"[OpenAI Vision OCR] Error: {e}")
        return ""

# ─── Modelos Pydantic ──────────────────────────────────────────────────────────
class CopilotoRequest(BaseModel):
    mensaje: str
    historial: list = []

class ScannerRequest(BaseModel):
    texto: str

# ─── Endpoint Copiloto legal ───────────────────────────────────────────────────
@app.post("/api/v1/copiloto", tags=["IA"])
async def copiloto_legal(req: CopilotoRequest):
    """Copiloto legal IA — powered by OpenAI gpt-4o-mini."""
    system_prompt = """Eres NEXO, un asistente legal especializado en Derecho de Familia chileno.

Tu conocimiento incluye:
- Ley N°14.908 (Alimentos)
- Código Civil (cuidado personal, relación directa y regular)
- Ley N°19.968 (Tribunales de Familia)
- Ley N°20.066 (Violencia Intrafamiliar)
- Mediación familiar obligatoria en Chile

IMPORTANTE:
- Habla en español, de manera clara y empática.
- Siempre aclara que tus respuestas son orientativas, no reemplazan a un abogado.
- Sé conciso: máximo 3 párrafos."""

    historial_texto = ""
    for msg in req.historial[-6:]:
        rol = "Usuario" if msg.get("role") == "user" else "NEXO"
        historial_texto += f"{rol}: {msg.get('content', '')}\n"

    prompt = f"{historial_texto}Usuario: {req.mensaje}\nNEXO:"

    try:
        respuesta = llamar_openai(prompt, system=system_prompt, temperatura=0.4)
        return {"respuesta": respuesta.strip(), "modelo": OPENAI_MODEL}
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Servicio de IA no disponible: {str(e)[:100]}")

# ─── Lógica Core del Scanner ────────────────────────────────────────────────
def _procesar_texto_scanner(texto: str, uid: str | None = None) -> dict:
    texto = texto.strip()
    if len(texto) < 20:
        raise HTTPException(status_code=400, detail="El texto o PDF es demasiado corto/ilegible.")

    # 💡 Recorte a 800 chars (–60% tokens)
    texto_seguro = pseudo_anonimizar(texto[:800])

    # 🔐 Límite diario por usuario
    if uid and not _verificar_limite_diario(uid):
        raise HTTPException(status_code=429, detail=f"Límite de {MAX_IA_REQUESTS_PER_DAY} análisis/día alcanzado.")

    # 🔥 CACHÉ MD5
    hash_key = _generar_hash_md5(texto_seguro)
    cached = _buscar_en_cache(hash_key)
    if cached is not None:
        cached["fuente"] = "cache"
        return cached

    system_prompt = (
        "Eres un guía práctico experto en Derecho de Familia chileno. "
        "Tu objetivo es explicar resoluciones a personas comunes sin conocimientos legales. "
        "PROHIBIDO usar lenguaje jurídico complejo, párrafos largos difíciles o asumir conocimientos legales. "
        "Responde SOLO con JSON válido, estructurado exactamente con estas 5 claves: "
        "'resumen_simple', 'significado', 'pasos', 'lugar', 'consejo'."
    )
    prompt = f"""Analiza esta resolución judicial chilena. Devuelve JSON estructurado así:
{{
  "resumen_simple": "1 o 2 líneas máximo, súper claro",
  "significado": "Qué significa esto en palabras muy simples",
  "pasos": [
    {{"paso": "1", "desc": "Acción concreta 1 (ej: ir a tribunal, reunir documentos)"}}
  ],
  "lugar": "Dónde hacerlo (comuna, tribunal en Chile si aplica)",
  "consejo": "Un consejo práctico para la persona"
}}

Resolución:
{texto_seguro}"""

    try:
        respuesta_raw = llamar_openai(prompt, system=system_prompt, temperatura=0.2)
        json_match = re.search(r'\{[\s\S]*\}', respuesta_raw)
        if not json_match:
            raise ValueError("La IA no devolvió JSON válido")

        analisis = json.loads(json_match.group(0))
        analisis["fuente"] = OPENAI_MODEL

        # 💾 Guardar en caché
        _guardar_en_cache(hash_key, texto_seguro, analisis)

        # 📝 Historial Firestore (solo metadata)
        if FIREBASE_ADMIN_OK and db is not None and uid:
            try:
                db.collection("usuarios").document(uid).collection("scanner_historial").document().set({
                    "resumen": analisis.get("resumen", ""),
                    "riesgo":  analisis.get("riesgo", "medio"),
                    "tipo":    analisis.get("tipo", "Desconocido"),
                    "timestamp": firestore.SERVER_TIMESTAMP,
                    "fuente": analisis["fuente"]
                })
            except Exception as e:
                logger.warning(f"⚠️ Error guardando historial: {e}")

        return analisis

    except Exception as e:
        logger.error(f"❌ [Scanner] Error: {e}")
        return {
            "resumen_simple": "El sistema de IA no pudo completar el análisis.",
            "significado": "Ocurrió un error temporal al intentar leer tu documento.",
            "pasos": [{"paso": "1", "desc": "Vuelve a intentarlo en unos minutos o contacta a soporte"}],
            "lugar": "Misma aplicación",
            "consejo": "Si el problema persiste, intenta subir una imagen o PDF más claro.",
            "fuente": "fallback",
            "error": str(e)[:100]
        }

# ─── Endpoint Scanner (Texto) ──────────────────────────────────────────────────
@app.post("/api/v1/scanner/analizar", tags=["IA"])
@limiter.limit("20/minute")
async def analizar_resolucion(request: Request, req: ScannerRequest, _user=Depends(verify_firebase_token)):
    """
    Analiza una resolución judicial con OpenAI gpt-4o-mini a partir de texto.
    🔥 Caché MD5 en Firestore activo.
    """
    uid = _user.get("uid") if _user else None
    return _procesar_texto_scanner(req.texto, uid)

# ─── Endpoint Scanner (Archivo PDF) ───────────────────────────────────────────
@app.post("/api/v1/scanner/subir", tags=["IA"])
@limiter.limit("20/minute")
async def subir_resolucion_scanner(request: Request, archivo: UploadFile = File(...), _user=Depends(verify_firebase_token)):
    """
    Analiza una resolución judicial subiendo un archivo PDF.
    Extrae texto y utiliza exactamente la misma lógica de IA y caché del scanner.
    """
    if not archivo.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Solo se permiten archivos PDF.")
    try:
        contenido = await archivo.read()
        doc = fitz.Document(stream=contenido, filetype="pdf")
        texto_crudo = "".join(p.get_text() for p in doc)
        
        uid = _user.get("uid") if _user else None
        return _procesar_texto_scanner(texto_crudo, uid)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ [Scanner PDF] Error: {e}")
        raise HTTPException(status_code=500, detail=f"Error procesando el PDF: {e}")

# ─── Subir sentencia colaborativa (PDF) ───────────────────────────────────────
@app.post("/api/v1/sentencias/subir", tags=["Colaborativo"])
async def subir_sentencia_colaborativa(archivo: UploadFile = File(...), _user=Depends(verify_firebase_token)):
    if not archivo.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Solo se permiten archivos PDF.")
    try:
        contenido = await archivo.read()
        doc = fitz.Document(stream=contenido, filetype="pdf")
        texto_crudo = "".join(p.get_text() for p in doc)
        if len(texto_crudo) < 10:
            raise HTTPException(status_code=422, detail="PDF sin texto extraíble (escaneo sin OCR).")

        texto_seguro = pseudo_anonimizar(texto_crudo[:800])

        system_prompt = "Eres un analista de datos judicial. Extrae información clave del fallo. Responde SOLO con JSON."
        prompt = f"""Lee esta sentencia y extrae:
{{"rit": "RIT o 'Desconocido'", "resultado": "Acoge|Rechaza|Acoge Parcialmente|Inadmisible", "monto_utm": 0, "resumen_anonimizado": "máximo 2 líneas sin nombres"}}

Sentencia:
{texto_seguro}"""

        respuesta = llamar_openai(prompt, system=system_prompt, temperatura=0.1)
        json_match = re.search(r'\{[\s\S]*?\}', respuesta)
        analisis = json.loads(json_match.group(0)) if json_match else {"rit": "DESCONOCIDO", "resultado": "Desconocido", "monto_utm": 0, "resumen_anonimizado": "Sin datos"}

        return {"mensaje": "Documento procesado.", "datos_extraidos": analisis, "preview": texto_seguro[:200]}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error procesando el documento: {e}")

# ─── Endpoints Analytics (SQLite) ─────────────────────────────────────────────
@app.get("/api/v1/analytics/velocidad", tags=["Analytics"])
def obtener_velocidad_tribunales():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT t.nombre AS tribunal,
               ROUND(AVG(JULIANDAY(c.fecha_sentencia) - JULIANDAY(SUBSTR(c.rit, -4) || '-01-01')) / 30.0, 1) AS meses
        FROM causas c JOIN tribunales t ON c.tribunal_id = t.id
        WHERE c.fecha_sentencia IS NOT NULL AND c.rit IS NOT NULL
        GROUP BY t.id HAVING COUNT(c.id) > 50
        ORDER BY meses ASC LIMIT 10
    """)
    resultados = [dict(row) for row in cursor.fetchall()]
    conn.close()
    for r in resultados:
        r["tribunal"] = r["tribunal"].replace("Juzgado de Familia ", "").replace("Juzgado de Letras y Garantía ", "").strip()
    return resultados

@app.get("/api/v1/analytics/estacionalidad", tags=["Analytics"])
def obtener_estacionalidad():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT strftime('%m', substr(fecha_sentencia, 1, 10)) as mes_num, COUNT(*) as sentencias
        FROM causas WHERE fecha_sentencia IS NOT NULL
        GROUP BY mes_num ORDER BY mes_num
    """)
    filas = cursor.fetchall()
    conn.close()
    meses = {'01':'Ene','02':'Feb','03':'Mar','04':'Abr','05':'May','06':'Jun','07':'Jul','08':'Ago','09':'Sep','10':'Oct','11':'Nov','12':'Dic'}
    return [{"mes": meses.get(f["mes_num"], "?"), "sentencias": f["sentencias"]} for f in filas if f["mes_num"] in meses]

@app.get("/api/v1/analytics/kpis", tags=["Analytics"])
def obtener_kpis_principales():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT t.nombre, ROUND(AVG(JULIANDAY(c.fecha_sentencia) - JULIANDAY(SUBSTR(c.rit, -4) || '-01-01')) / 30.0, 1) AS meses
        FROM causas c JOIN tribunales t ON c.tribunal_id = t.id
        WHERE c.fecha_sentencia IS NOT NULL AND c.rit IS NOT NULL
        GROUP BY t.id HAVING COUNT(c.id) > 100 ORDER BY meses ASC LIMIT 1
    """)
    res_velocidad = cursor.fetchone()
    cursor.execute("""
        SELECT j.nombre_completo, COUNT(cj.causa_id) as total_fallos
        FROM jueces j JOIN causa_juez cj ON j.id = cj.juez_id
        GROUP BY j.id ORDER BY total_fallos DESC LIMIT 1
    """)
    res_carga = cursor.fetchone()
    cursor.execute("""
        SELECT strftime('%m', substr(fecha_sentencia, 1, 10)) as mes_num, COUNT(*) as total
        FROM causas WHERE fecha_sentencia IS NOT NULL
        GROUP BY mes_num ORDER BY total DESC LIMIT 1
    """)
    res_estacional = cursor.fetchone()
    conn.close()
    meses_map = {'01':'Enero','02':'Febrero','03':'Marzo','04':'Abril','05':'Mayo','06':'Junio','07':'Julio','08':'Agosto','09':'Septiembre','10':'Octubre','11':'Noviembre','12':'Diciembre'}
    return {
        "velocidad": {"valor": f"{res_velocidad['meses']} meses" if res_velocidad else "N/A", "subtitulo": res_velocidad['nombre'].replace("Juzgado de Familia ", "") if res_velocidad else "Sin datos"},
        "carga":     {"valor": f"{res_carga['total_fallos']} fallos" if res_carga else "N/A", "subtitulo": res_carga['nombre_completo'] if res_carga else "Sin datos"},
        "estacional": {"valor": f"{res_estacional['total']} fallos" if res_estacional else "N/A", "subtitulo": f"Mes de {meses_map.get(res_estacional['mes_num'], '?')}" if res_estacional else "Sin datos"}
    }

@app.get("/api/v1/analytics/materias", tags=["Analytics"])
def obtener_distribucion_materias():
    conn = get_db_connection()
    cursor = conn.cursor()
    cursor.execute("""
        SELECT materia as nombre, COUNT(*) as valor
        FROM causas WHERE materia IS NOT NULL AND materia != ''
        GROUP BY materia ORDER BY valor DESC LIMIT 5
    """)
    resultados = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return resultados

# =====================================================================
#  NUEVO MÓDULO: EXTRACCIÓN ESTRUCTURADA DE CAUSAS 
# =====================================================================

async def _procesar_causa_unificada(texto_seguro: str, origen: str, uid: str):
    import hashlib
    # 1. MD5
    md5_hash = hashlib.md5(texto_seguro.encode('utf-8')).hexdigest()
    
    # 2. Check duplicate in Firestore for this user
    if db:
        try:
            docs = db.collection('causas').where('userId', '==', uid).where('hashMD5', '==', md5_hash).limit(1).get()
            if docs:
                logger.info(f"⚡ [Causas] Cache hit MD5: {md5_hash} para user {uid}")
                return docs[0].to_dict()
        except Exception as db_err:
            logger.error(f"⚠️ Error leyendo caché MD5 de Firebase: {db_err}")

    # 3. Rate Limit IA check
    if not _verificar_limite_diario(uid):
        return {"error": "Límite de análisis de IA diario alcanzado. Intenta mañana."}
    
    # 4. Prompt
    system_prompt = (
        "Eres un software extractor de datos judiciales de Chile. "
        "Solo debes retornar un objeto JSON, sin agregar texto plano de formato markdown alrededor. "
        "Si no encuentras un dato, establece la clave en null de manera estricta."
    )
    prompt = f"""Analiza el siguiente texto (puede venir de un PDF, imagen o texto ingresado por usuario) y extrae información legal estructurada de Chile.

Responde SOLO en JSON válido.

Si no encuentras un dato, usa null.

Incluye un resumen_simple explicado en lenguaje fácil para cualquier persona (fuera del objeto 'partes').

Debes usar exactamente esta estructura de JSON con estas claves:
{{
"rit": "",
"ruc": "",
"tribunal": "",
"juez": "",
"tipo_causa": "",
"estado_causa": "",
"fecha": "",
"partes": {{
  "demandante": "",
  "demandado": ""
}},
"resumen_simple": ""
}}

Texto:
{texto_seguro}"""

    # 5. Call OpenAI con más tokens permitidos
    respuesta_raw = llamar_openai(prompt, system=system_prompt, temperatura=0.1, max_tokens=1500)
    
    # 6. Parse JSON
    try:
        json_match = re.search(r'\{[\s\S]*\}', respuesta_raw)
        if not json_match:
            raise ValueError(f"No se encontró JSON válido. Respondió: {respuesta_raw[:200]}")
        causa_datos = json.loads(json_match.group(0))
    except Exception as e:
        logger.error(f"Error parseando JSON de causa: {e}")
        causa_datos = {
             "rit": None, "ruc": None, "tribunal": None, "juez": None,
             "tipo_causa": None, "estado_causa": None, "fecha": None,
             "partes": {"demandante": None, "demandado": None},
             "resumen_simple": "Error IA estructurando"
        }
    
    # Validar formato interno de partes por las dudas (a veces la IA mete todo suelto)
    partes = causa_datos.get("partes")
    if not isinstance(partes, dict):
        partes = {"demandante": causa_datos.get("demandante"), "demandado": causa_datos.get("demandado")}
    
    # 7. Structure final doc
    doc_final = {
        "userId": uid,
        "rit": causa_datos.get("rit"),
        "ruc": causa_datos.get("ruc"),
        "tribunal": causa_datos.get("tribunal"),
        "juez": causa_datos.get("juez"),
        "tipoCausa": causa_datos.get("tipo_causa"),
        "estadoCausa": causa_datos.get("estado_causa"),
        "fecha": causa_datos.get("fecha"),
        "partes": partes,
        "resumenIA": causa_datos.get("resumen_simple"),
        "hashMD5": md5_hash,
        "createdAt": datetime.now(timezone.utc).isoformat(),
        "origen": "pdf"
    }
    
    # 8. Guardar en Firestore
    if db:
        try:
            db.collection('causas').add(doc_final)
            logger.info("✅ Documento guardado correctamente en Firebase")
        except Exception as e:
            logger.error(f"Error guardando causa en FB: {e}")

    return doc_final

@app.post("/api/v1/causas/procesar")
async def procesar_causa_endpoint(
    request: Request,
    archivo: UploadFile = File(None),
    texto: str = Form(None),
    usuario: dict = Depends(verify_firebase_token)
):
    if not usuario:
        raise HTTPException(status_code=401, detail="Usuario no autenticado.")
    
    uid = str(usuario.get("uid", ""))
    
    texto_extraido = ""
    origen_doc = "texto"

    if archivo and archivo.filename:
        contenido = await archivo.read()
        filename_lower = archivo.filename.lower()
        if filename_lower.endswith('.pdf'):
            origen_doc = "pdf"
            try:
                doc = fitz.open(stream=contenido, filetype="pdf")
                for pagina in doc:
                    texto_extraido += pagina.get_text("text") + "\n"
                doc.close()
            except Exception as e:
                logger.error(f"Error leyendo PDF causa: {e}")
                raise HTTPException(status_code=400, detail="Error leyendo el archivo PDF.")
        elif filename_lower.endswith(('.png', '.jpg', '.jpeg', '.webp')):
            origen_doc = "imagen"
            try:
                b64_img = base64.b64encode(contenido).decode('utf-8')
                texto_extraido = llamar_openai_vision(b64_img)
            except Exception as e:
                logger.error(f"Error OCR Visión: {e}")
                raise HTTPException(status_code=400, detail="Error procesando la imagen por OCR.")
        else:
            raise HTTPException(status_code=400, detail="Formato de archivo no soportado.")
    elif texto:
        origen_doc = "texto"
        texto_extraido = texto
    else:
        raise HTTPException(status_code=400, detail="Debes enviar un archivo o texto.")
        
    texto_limpio = texto_extraido.strip()
    if len(texto_limpio) < 20:
        raise HTTPException(status_code=400, detail="El material provisto no contiene suficiente texto legible.")
        
    logger.info(f"⚡ [Causas] Procesando - Origen: {origen_doc}")
    texto_seguro = re.sub(r'[^\w\s.,;:()\-/\'"áéíóúÁÉÍÓÚñÑ]', '', texto_limpio[:25000])
    
    resultado = await _procesar_causa_unificada(texto_seguro, origen_doc, uid)
    
    logger.info(f"📊 [Result Logs] Tipo input: {origen_doc} | Longitud texto procesado: {len(texto_seguro)}")
    return resultado

@app.get("/api/v1/causas/mis-causas")
def mis_causas_endpoint(request: Request, usuario: dict = Depends(verify_firebase_token)):
    if not usuario:
        raise HTTPException(status_code=401, detail="Usuario no autenticado.")
    
    uid = str(usuario.get("uid", ""))
    if not db:
        raise HTTPException(status_code=500, detail="Base de datos no disponible.")
        
    try:
        docs = db.collection("causas").where('userId', '==', uid).get()
        resultados = [doc.to_dict() for doc in docs]
        # Ordenar por fecha descendente usando Python para evitar requerir index compuesto en FB
        resultados.sort(key=lambda x: x.get('createdAt', ''), reverse=True)
        return resultados
    except Exception as e:
        logger.error(f"Error obteniendo causas: {e}")
        return []


# =====================================================================
#  MÓDULO: GENERADOR DE ESCRITOS LEGALES PROFESIONALES
# =====================================================================

TIPOS_ESCRITO = {
    "rebaja_pension": "Solicitud de Rebaja de Pensión Alimenticia",
    "cumplimiento_pension": "Solicitud de Cumplimiento de Pensión Alimenticia",
    "solicitud_liquidacion": "Solicitud de Liquidación de Deuda de Alimentos",
    "cese_alimentos": "Solicitud de Cese de Obligación Alimenticia",
    "regimen_visitas": "Solicitud de Régimen de Relación Directa y Regular",
    "medidas_apremio": "Solicitud de Medidas de Apremio por Incumplimiento",
}

LEYES_REFERENCIA = {
    "rebaja_pension": ["Ley N°14.908 art. 10 (rebaja por cambio de circunstancias)", "Código Civil art. 332"],
    "cumplimiento_pension": ["Ley N°14.908 art. 14 (apremios)", "Ley N°19.968 art. 66 (ejecución de sentencias)"],
    "solicitud_liquidacion": ["Ley N°14.908 art. 18 (liquidación de deuda)", "Auto Acordado CS sobre liquidaciones"],
    "cese_alimentos": ["Código Civil art. 332 (causales de cese)", "Ley N°14.908 art. 10 inc. final"],
    "regimen_visitas": ["Ley N°19.968 art. 48 (acuerdos sobre relación directa)", "Código Civil art. 229 (derecho del hijo)"],
    "medidas_apremio": ["Ley N°14.908 art. 14 (arresto nocturno)", "Ley N°14.908 art. 16 (retención de fondos)"],
}

class EscritoRequest(BaseModel):
    tipo_escrito: str          # uno de TIPOS_ESCRITO keys
    situacion: str             # descripción libre del usuario
    tribunal: str = ""
    rit: str = ""
    nombre_usuario: str = "COMPARECIENTE"
    rut_usuario: str = ""
    direccion_usuario: str = ""
    telefono_usuario: str = ""
    email_usuario: str = ""
    contraparte: str = ""

@app.post("/api/v1/escritos/generar", tags=["Escritos"])
@limiter.limit("10/minute")
async def generar_escrito(
    request: Request,
    req: EscritoRequest,
    _user=Depends(verify_firebase_token)
):
    """
    Genera un escrito legal profesional al nivel de un abogado chileno.
    Produce: (A) Escrito formal para tribunal + (B) Explicación simple para el usuario.
    """
    if req.rut_usuario and req.rut_usuario.strip() != "":
        if not validar_rut_chileno(req.rut_usuario):
            raise HTTPException(status_code=400, detail="El RUT ingresado no es válido. Verifica el dígito verificador y formato.")

    uid = _user.get("uid") if _user else "test-user"

    tipo_label = TIPOS_ESCRITO.get(req.tipo_escrito, req.tipo_escrito)
    leyes = LEYES_REFERENCIA.get(req.tipo_escrito, ["Ley N°14.908", "Ley N°19.968"])
    leyes_texto = " | ".join(leyes)

    tribunal_str = req.tribunal or "TRIBUNAL DE FAMILIA COMPETENTE"
    rit_str = f"RIT: {req.rit}" if req.rit else "RIT: A determinarse"
    nombre = req.nombre_usuario or "EL/LA COMPARECIENTE"
    rut = f", RUT {req.rut_usuario}," if req.rut_usuario else ","
    direccion = f" domiciliado en {req.direccion_usuario}," if req.direccion_usuario else " domiciliado en [COMPLETAR DOMICILIO],"
    
    # Datos de contacto para Otrosí
    contacto = []
    if req.telefono_usuario: contacto.append(f"teléfono {req.telefono_usuario}")
    if req.email_usuario: contacto.append(f"correo electrónico {req.email_usuario}")
    contacto_str = " y ".join(contacto) if contacto else "[COMPLETAR DATOS DE CONTACTO]"
    
    contraparte = req.contraparte or "la contraparte"

    system_prompt = """Eres un abogado senior especialista en Derecho de Familia chileno con 20 años de experiencia litigando en Tribunales de Familia.

Tu rol es redactar escritos judiciales que:
1. Sean formalmente correctos según la práctica forense chilena
2. Usen terminología jurídica precisa pero accesible
3. Citen leyes chilenas vigentes y aplicables al caso concreto
4. Tengan la estructura EXACTA que exigen los Juzgados de Familia
5. Sean presentables en tribunal SIN modificaciones importantes

REGLAS ABSOLUTAS:
- NUNCA incluyas explicaciones al usuario dentro del escrito
- NUNCA uses frases como "te recomiendo", "deberías", "mi consejo"
- SIEMPRE usa "S.S." o "S.S.ª" para referirte al juez
- SIEMPRE usa "SS.MM." para el plural
- El escrito debe sonar como redactado por un abogado
- Usa "Por tanto" o "Por estas razones" antes de la petición principal
- La petición debe ser CONCRETA y ESPECÍFICA
- ROL ESTRICTO DE PARTES: El demandante SIEMPRE es el usuario compareciente ({req.nombre_usuario}) y la contraparte es el beneficiario (hijo/a) representado por el otro progenitor. NUNCA pongas al usuario como demandante y demandado a la vez.

ESTRUCTURA OBLIGATORIA DEL ESCRITO:
1. **SUMA**: Precisa y al punto (ej: "SOLICITA REBAJA DE PENSIÓN ALIMENTICIA.")
2. **INDIVIDUALIZACIÓN**: NUNCA digas "a favor de mi contraparte". Debes decir: "en favor de mi hijo/a [Nombre Completo Hijo/a], representado legalmente por [Nombre de la madre/padre o contraparte]". Si no tienes el nombre del hijo, usa "[NOMBRE DEL HIJO/A]".
3. **HECHOS**: Si es por cesantía, incluye fechas exactas (o marcadores [Fecha]) de despido y menciona explícitamente que el vínculo terminó "según consta en finiquito que se acompaña".
4. **DERECHO**: Cita los artículos pertinentes.
5. **PRUEBA**: Siempre incluye una sección explícita de documentos a acompañar. OBLIGATORIO: Para demandas de familia como rebaja, añade "1. Acta de mediación frustrada (Obligatorio por ley). 2. Certificado de nacimiento. 3. Finiquito de trabajo (si aplica)."
6. **OTROSÍ DE NOTIFICACIONES**: Menciona la aceptación expresa de notificaciones por correo electrónico.
7. **OTROSÍ DE PATROCINIO Y PODER**: SIEMPRE añade un Otrosí designando abogado patrocinante y apoderado ("Vengo en designar abogado patrocinante y conferir poder a...").

Responde SOLO con un JSON con exactamente estas claves:
{
  "escrito_formal": "texto completo del escrito con el formato legal chileno descrito",
  "explicacion_simple": "explicación en 3-4 párrafos cortos para el usuario sin términos legales",
  "advertencias": ["lista de cosas urgentes, EJ: recordar que en Chile los tribunales evalúan activos en AFP o ahorros antes de conceder rebajas solo por cesantía"]
}"""

    prompt = f"""Redacta un escrito judicial de tipo: {tipo_label}

DATOS DEL CASO:
- Tribunal: {tribunal_str}
- {rit_str}
- Nombre del compareciente: {nombre}
- RUT: {req.rut_usuario or "[COMPLETAR RUT]"}
- Dirección: {req.direccion_usuario or "[COMPLETAR DIRECCIÓN]"}
- Contacto: {contacto_str}
- Representante legal de los menores (contraparte): {contraparte or "[NOMBRE CONTRAPARTE]"}
- Situación / Hechos del usuario: {req.situacion}

LEYES APLICABLES A CITAR: {leyes_texto}

ESTRUCTURA OBLIGATORIA DEL ESCRITO:
1. ENCABEZADO: "{tribunal_str} / {rit_str}"
2. COMPARECENCIA: "Yo, {nombre}{rut}{direccion} en autos caratulados [NOMBRE CAUSA]..."
3. EXPONGO / EXPONE: Hechos ordenados numéricamente, claros y concisos, basados en la "Situación del usuario"
4. FUNDAMENTOS DE DERECHO: Citar artículos exactos de las leyes {leyes_texto}
5. POR TANTO: Petición concreta al tribunal
6. OTROSÍ (Forma de Notificación): Solicitar que las notificaciones se realicen al {contacto_str}. Agrega otros otrosíes si el tipo de escrito lo requiere.
7. [Ciudad], [fecha] / Firma: "{nombre} / RUT: {req.rut_usuario or '___________'}"

IMPORTANTE: El escrito debe poder presentarse directamente en tribunal. Usa [COMPLETAR] para campos que el usuario debe llenar (domicilio, fechas específicas, montos concretos si no se proporcionaron).

Genera el escrito ahora:"""

    try:
        respuesta_raw = llamar_openai(prompt, system=system_prompt, temperatura=0.15, max_tokens=2000)
        json_match = re.search(r'\{[\s\S]*\}', respuesta_raw)
        if not json_match:
            raise ValueError("La IA no devolvió JSON válido")
        resultado = json.loads(json_match.group(0))
        
        escrito_formal_text = resultado.get("escrito_formal", "")
        pdf_base64 = generar_pdf_basico(escrito_formal_text, req)
        docx_base64 = generar_docx_basico(escrito_formal_text, req)
        
        # Hardcoded warnings para rebaja_pension
        advertencias = resultado.get("advertencias", [])
        if req.tipo_escrito == "rebaja_pension":
            if not any("mediación" in str(x).lower() for x in advertencias):
                advertencias.insert(0, "¡IMPORTANTE! En Chile, para demandar rebaja de alimentos, es REQUISITO DE ADMISIBILIDAD contar con el Certificado de Mediación Frustrada. Si no lo tienes, la demanda será rechazada de plano.")
        
    except Exception as e:
        logger.error(f"[Escritos] Error generando escrito: {e}")
        raise HTTPException(status_code=503, detail=f"Error generando el escrito: {str(e)[:120]}")

    # Guardar en Firestore para historial
    if db and uid:
        try:
            db.collection("escritos").add({
                "userId": uid,
                "tipo": req.tipo_escrito,
                "tipoLabel": tipo_label,
                "tribunal": req.tribunal,
                "rit": req.rit,
                "datos_personales": {
                    "nombre": req.nombre_usuario,
                    "rut": req.rut_usuario,
                    "direccion": req.direccion_usuario,
                    "telefono": req.telefono_usuario,
                    "email": req.email_usuario
                },
                "createdAt": datetime.now(timezone.utc).isoformat(),
                "preview": resultado.get("escrito_formal", "")[:200],
            })
            logger.info(f"✅ Escrito guardado en Firebase para user {uid}")
        except Exception as e:
            logger.warning(f"[Escritos] No se pudo guardar en Firebase: {e}")

    return {
        "tipo": req.tipo_escrito,
        "tipo_label": tipo_label,
        "escrito_formal": escrito_formal_text,
        "explicacion_simple": resultado.get("explicacion_simple", ""),
        "advertencias": advertencias,
        "pdf_base64": pdf_base64,
        "docx_base64": docx_base64,
        "leyes_citadas": leyes,
    }

def generar_pdf_basico(texto_escrito: str, req: EscritoRequest) -> str:
    if not REPORTLAB_OK: return ""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=letter,
        rightMargin=2.0*cm, leftMargin=3.0*cm, topMargin=2.5*cm, bottomMargin=2.5*cm
    )
    styles = getSampleStyleSheet()
    style_normal = ParagraphStyle(
        'LegalNormal', parent=styles['Normal'],
        fontName='Helvetica', fontSize=12, leading=18, alignment=TA_JUSTIFY
    )
    style_header = ParagraphStyle(
        'LegalHeader', parent=styles['Normal'],
        fontName='Helvetica-Bold', fontSize=12, leading=15, alignment=TA_RIGHT
    )
    style_center = ParagraphStyle(name='Center', fontName='Helvetica', fontSize=12, alignment=TA_CENTER)
    style_small = ParagraphStyle(name='Small', fontName='Helvetica', fontSize=8, alignment=TA_CENTER)
    
    story = []
    lines = texto_escrito.split('\n')
    body_started = False
    
    for line in lines:
        text = line.strip()
        if not text:
            story.append(Spacer(1, 12))
            continue
            
        if not body_started and any(x in text.upper() for x in ["SUMA:", "S U M A :", "RIT:", "MATERIA:", "EN LO PRINCIPAL", "SOLICITA:"]):
            story.append(Paragraph(text, style_header))
            continue
            
        if "S.J.L." in text.upper() or "S.S." in text.upper() or "S.S.ª" in text.upper():
            body_started = True
            story.append(Spacer(1, 12))
            
        if text.upper().startswith("OTROSÍ"):
            story.append(Paragraph(f"<b>{text}</b>", style_normal))
        else:
            story.append(Paragraph(text, style_normal))
            
        story.append(Spacer(1, 6))
            
    # Firma
    story.append(Spacer(1, 3*cm))
    story.append(Paragraph("___________________________", style_center))
    story.append(Paragraph(req.nombre_usuario.upper(), style_center))
    story.append(Paragraph(f"RUT: {req.rut_usuario}", style_center))
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph("Documento generado para ser firmado mediante Clave Única en la Oficina Judicial Virtual.", style_small))

    try:
        doc.build(story)
        return base64.b64encode(buffer.getvalue()).decode('utf-8')
    except Exception as e:
        logger.error(f"[ReportLab] Error: {e}")
        return ""

def generar_docx_basico(texto_escrito: str, req: EscritoRequest) -> str:
    if not DOCX_OK: return ""
    try:
        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(3.0)
            section.right_margin = Cm(2.0)
            
        lines = texto_escrito.split('\n')
        body_started = False
        
        for line in lines:
            text = line.strip()
            if not text:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run(text)
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            
            if not body_started and any(x in text.upper() for x in ["SUMA:", "S U M A :", "RIT:", "MATERIA:", "EN LO PRINCIPAL", "SOLICITA:"]):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run.bold = True
                continue
                
            if "S.J.L." in text.upper() or "S.S." in text.upper() or "S.S.ª" in text.upper():
                body_started = True
                
            if text.upper().startswith("OTROSÍ"):
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_paragraph()
        p_firma = doc.add_paragraph()
        p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_f = p_firma.add_run(f"___________________________\n{req.nombre_usuario.upper()}\n[RUT: {req.rut_usuario}]")
        run_f.font.name = 'Arial'
        run_f.font.size = Pt(12)
        
        p_nota = doc.add_paragraph()
        p_nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_nota = p_nota.add_run("\nDocumento generado para ser firmado mediante Clave Única en la Oficina Judicial Virtual.")
        run_nota.font.name = 'Arial'
        run_nota.font.size = Pt(8)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        return base64.b64encode(buffer.getvalue()).decode('utf-8')
    except Exception as e:
        logger.error(f"[python-docx] Error: {e}")
        return ""    
@app.get("/api/v1/escritos/tipos", tags=["Escritos"])
async def listar_tipos_escrito():
    """Lista todos los tipos de escritos disponibles."""

